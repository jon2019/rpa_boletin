from pathlib import Path
import sys

ROOT_DIR = Path(__file__).resolve().parents[2]
SRC_DIR = ROOT_DIR / 'src'
if str(SRC_DIR) not in sys.path:
    sys.path.insert(0, str(SRC_DIR))

"""Script para generar documentación consolidada en DOCX desde docs/."""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

DOCS_DIR = ROOT_DIR / 'docs'
OUTPUT_DIR = DOCS_DIR / 'legacy' / 'ejecutiva'
OUTPUT_PATH = OUTPUT_DIR / 'documentacion_completa.docx'

SECCIONES = [
    (DOCS_DIR / 'README.md', '1. Documentación General'),
    (DOCS_DIR / 'operacion' / 'README_sistema.md', '2. Operación del Sistema'),
    (DOCS_DIR / 'datos' / 'modelo_datos.md', '3. Modelo de Datos'),
    (DOCS_DIR / 'datos' / 'modelo_relacional.md', '4. Modelo Relacional'),
    (DOCS_DIR / 'negocio' / 'reglas_negocio.md', '5. Reglas de Negocio'),
]


def create_documentation_docx() -> Path:
    """Genera el archivo de documentación completa en formato DOCX."""
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    styles = doc.styles

    title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.size = Pt(24)
    title_style.font.bold = True
    title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    heading1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    heading1_style.font.size = Pt(18)
    heading1_style.font.bold = True

    heading2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
    heading2_style.font.size = Pt(16)
    heading2_style.font.bold = True

    heading3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
    heading3_style.font.size = Pt(14)
    heading3_style.font.bold = True

    normal_style = styles.add_style('CustomNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_style.font.size = Pt(11)

    code_style = styles.add_style('CustomCode', WD_STYLE_TYPE.PARAGRAPH)
    code_style.font.name = 'Courier New'
    code_style.font.size = Pt(10)

    title = doc.add_paragraph('Documentación del Sistema', style='CustomTitle')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle = doc.add_paragraph('RPA Boletín Minero-Energético', style='CustomTitle')
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('Fuente: documentación canónica del directorio docs/', style='CustomNormal')
    doc.add_paragraph('Fecha: Abril 2026', style='CustomNormal')
    doc.add_page_break()

    doc.add_paragraph('Índice', style='CustomHeading1')
    for _, titulo in SECCIONES:
        doc.add_paragraph(titulo, style='CustomNormal')
    doc.add_page_break()

    def process_markdown_file(filepath: Path, section_title: str) -> None:
        if not filepath.exists():
            doc.add_paragraph(f'Archivo no encontrado: {filepath}', style='CustomNormal')
            doc.add_page_break()
            return

        doc.add_paragraph(section_title, style='CustomHeading1')
        content = filepath.read_text(encoding='utf-8')
        lines = content.splitlines()
        in_code_block = False

        for raw_line in lines:
            line = raw_line.strip()
            if not line:
                continue
            if line.startswith('# '):
                doc.add_paragraph(line[2:], style='CustomHeading1')
            elif line.startswith('## '):
                doc.add_paragraph(line[3:], style='CustomHeading2')
            elif line.startswith('### '):
                doc.add_paragraph(line[4:], style='CustomHeading3')
            elif line.startswith('```'):
                in_code_block = not in_code_block
            elif in_code_block:
                doc.add_paragraph(line, style='CustomCode')
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(f'• {line[2:]}', style='CustomNormal')
            elif '|' in line and ('Campo' in line or 'Tabla' in line or 'Columna' in line):
                doc.add_paragraph(line.replace('|', ' | '), style='CustomCode')
            else:
                doc.add_paragraph(line, style='CustomNormal')

        doc.add_page_break()

    for filepath, section_title in SECCIONES:
        process_markdown_file(filepath, section_title)

    doc.save(OUTPUT_PATH)
    print(f'Documentación generada: {OUTPUT_PATH}')
    return OUTPUT_PATH


if __name__ == '__main__':
    create_documentation_docx()
